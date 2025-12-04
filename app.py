import os
import re
import time
import shutil
import zipfile
import io

from flask import (
    Flask,
    request,
    jsonify,
    render_template,
    send_from_directory,
    send_file,
    url_for,
)
from werkzeug.utils import secure_filename
from bs4 import BeautifulSoup
from docx import Document

# -------------------------------------------------------
# Environment-aware upload folder
# -------------------------------------------------------
RUNNING_FLY = os.environ.get("FLY_APP_NAME") is not None

if RUNNING_FLY:
    # On Fly.io: use the mounted volume
    UPLOAD_FOLDER = "/data/uploads"
else:
    # Local dev: ./uploads
    UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {"html", "zip"}

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

EXPECTED_BANNER_SIZES = [
    (300, 250),
    (160, 600),
    (728, 90),
    (320, 50),
    (970, 250),
    (300, 50),
    (300, 600),
]

MAX_FILE_SIZE_KB = 150
ANIMATION_MAX_DURATION = 15
MAX_LOOP_COUNT = 3
ANIMATION_WARNING_THRESHOLD = 14.8


# -------------------------------------------------------
# AUTO CLEANUP OF OLD UPLOADS
# -------------------------------------------------------
def cleanup_old_uploads(max_age_hours: int = 6) -> None:
    """Delete files/folders in UPLOAD_FOLDER older than max_age_hours."""
    now = time.time()
    max_age_seconds = max_age_hours * 3600

    for item in os.listdir(UPLOAD_FOLDER):
        path = os.path.join(UPLOAD_FOLDER, item)
        try:
            mtime = os.path.getmtime(path)
        except FileNotFoundError:
            continue

        if now - mtime > max_age_seconds:
            if os.path.isdir(path):
                shutil.rmtree(path, ignore_errors=True)
            else:
                os.remove(path)
            print(f"[CLEANUP] Deleted old upload: {path}")


# -------------------------------------------------------
# Helper Functions
# -------------------------------------------------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_zip(file_path: str) -> str | None:
    """Extract a ZIP into UPLOAD_FOLDER/<zip_name> and return that folder path."""
    extracted_folder = os.path.join(
        UPLOAD_FOLDER, os.path.splitext(os.path.basename(file_path))[0]
    )
    try:
        with zipfile.ZipFile(file_path, "r") as zip_ref:
            zip_ref.extractall(extracted_folder)
    except zipfile.BadZipFile:
        return None
    return extracted_folder


def check_border_in_css(css_content: str) -> bool:
    return bool(re.search(r"border\s*:\s*1px\s+solid", css_content, re.IGNORECASE))


def extract_size_from_inline(style: str) -> tuple[int | None, int | None]:
    w = re.search(r"width\s*:\s*(\d+)px", style)
    h = re.search(r"height\s*:\s*(\d+)px", style)
    return (int(w.group(1)), int(h.group(1))) if w and h else (None, None)


def infer_size_from_filename(path: str) -> tuple[int | None, int | None]:
    """
    Try to infer size from filename like ..._300x250_... or banner_160x600.html
    """
    name = os.path.basename(path)
    match = re.search(r"(\d{2,4})x(\d{2,4})", name)
    if match:
        return int(match.group(1)), int(match.group(2))
    return None, None


def extract_size_from_css_text(css_text: str) -> tuple[int | None, int | None]:
    """
    Try to extract width/height specifically from rules that look like
    .adSize, #viewport, #banner, #stage, etc., instead of grabbing
    random 16px heights from small elements.
    """
    # Only look at likely container selectors
    selector_patterns = [
        r"\.adSize\b",
        r"#viewport\b",
        r"#banner\b",
        r"#ad\b",
        r"#stage\b",
        r"#designContainer\b",
    ]

    for sel in selector_patterns:
        # Capture the block for that selector: selector { ... }
        pattern = rf"{sel}[^\{{]*\{{([^}}]+)\}}"
        m = re.search(pattern, css_text, re.IGNORECASE | re.DOTALL)
        if not m:
            continue

        block = m.group(1)
        w = re.search(r"width\s*:\s*(\d+)px", block)
        h = re.search(r"height\s*:\s*(\d+)px", block)
        if w and h:
            return int(w.group(1)), int(h.group(1))

    # If nothing found, don't guess from random width/height in the file.
    return None, None


def detect_banner_size(
    soup: BeautifulSoup, css_paths: list[str], html_file_path: str
) -> tuple[int | None, int | None, str]:
    """
    Heuristic for banner size with safe priority:

    1) META ad.size
    2) Inline styles on .adSize or known root IDs (viewport, banner, etc.)
    3) CSS rules for .adSize / #viewport / etc.
    4) File name like _300x250
    """
    width = height = None
    source = "Unknown"

    # 1) META TAG (highest priority)
    meta = soup.find("meta", {"name": "ad.size"})
    if meta and meta.get("content"):
        match = re.search(r"width=(\d+),\s*height=(\d+)", meta["content"])
        if match:
            width = int(match.group(1))
            height = int(match.group(2))
            return width, height, "META ad.size"

    # 2) Inline on .adSize or root containers
    root_ids = {"viewport", "banner", "ad", "stage", "designContainer"}

    candidates = []

    # Any element with class containing "adSize"
    for tag in soup.find_all(True):
        classes = tag.get("class") or []
        if any(c == "adSize" for c in classes):
            candidates.append(tag)

    # Known root IDs
    for rid in root_ids:
        t = soup.find(id=rid)
        if t and t not in candidates:
            candidates.append(t)

    for tag in candidates:
        style = tag.get("style", "")
        w, h = extract_size_from_inline(style)
        if w and h:
            return w, h, "Inline root (.adSize / root id)"

    # 3) CSS rules from linked stylesheets
    for css_path in css_paths:
        if not os.path.exists(css_path):
            continue
        try:
            with open(css_path, "r", encoding="utf-8", errors="replace") as f:
                css_text = f.read()
        except OSError:
            continue

        w, h = extract_size_from_css_text(css_text)
        if w and h:
            return w, h, f"CSS ({os.path.basename(css_path)})"

    # 4) Filename fallback
    w, h = infer_size_from_filename(html_file_path)
    if w and h:
        return w, h, "Filename pattern"

    return None, None, source


# -------------------------------------------------------
#  HTML VALIDATION
# -------------------------------------------------------
def validate_html(file_path: str) -> dict:
    results = {
        "warnings": [],
        "errors": [],
        "banner_size": None,
        "border": "❌ Missing 1px border",
        "used_meta_tag": False,
        "size_source": "Unknown",
    }

    if not os.path.exists(file_path):
        results["errors"].append(f"File not found: {file_path}")
        return results

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f, "html.parser")

    base_path = os.path.dirname(file_path)
    missing_assets: list[str] = []
    border_found = False
    border_source = None

    css_paths: list[str] = []

    # ------------------------ CSS Parsing ------------------------
    for link_tag in soup.find_all("link", rel="stylesheet"):
        href = link_tag.get("href")
        if not href:
            continue
        if href.startswith(("http://", "https://")):
            # external CSS – skip size/border detection here
            continue
        css_path = os.path.join(base_path, href.lstrip("/"))
        css_paths.append(css_path)

        if not os.path.exists(css_path):
            continue

        try:
            with open(css_path, "r", encoding="utf-8", errors="replace") as f:
                css_text = f.read()
        except OSError:
            continue

        # Border detection in CSS
        if check_border_in_css(css_text):
            border_found = True
            border_source = "CSS"

    # ------------------------ Detect size using safe heuristic ------------------------
    bw, bh, size_source = detect_banner_size(soup, css_paths, file_path)
    if bw and bh:
        results["banner_size"] = f"{bw}x{bh}"
        results["size_source"] = size_source
        # Only treat as error if clearly not one of the standard expected sizes
        if (bw, bh) not in EXPECTED_BANNER_SIZES:
            results["errors"].append(
                f"Banner size {bw}x{bh} is not in expected list. Please double-check spec."
            )
    else:
        results["warnings"].append(
            "⚠️ Could not confidently determine banner size. Please double-check."
        )

    # META tag flag
    meta = soup.find("meta", {"name": "ad.size"})
    if meta and meta.get("content"):
        if re.search(r"width=\d+,\s*height=\d+", meta["content"]):
            results["used_meta_tag"] = True

    # ------------------------ Inline border ------------------------
    for tag in soup.find_all(True):
        style = tag.get("style", "")
        if (
            "border" in style.lower()
            and "1px" in style
            and "solid" in style.lower()
        ):
            border_found = True
            border_source = "Inline Style"
            break

    # ------------------------ SVG border ------------------------
    for rect in soup.find_all("rect"):
        stroke = rect.get("stroke")
        stroke_w = rect.get("stroke-width")
        try:
            if stroke and stroke_w is not None and float(stroke_w) == 1:
                border_found = True
                border_source = "SVG Border"
                break
        except Exception:
            continue

    # ------------------------ Final Border ------------------------
    if border_found:
        results["border"] = f"✅ 1px border present ({border_source})"
    else:
        # Still flexible: this is a warning-tier issue you can interpret
        results["border"] = "❌ Missing 1px border (or not detected)"

    # ------------------------ Missing Assets ------------------------
    for tag in soup.find_all(["img", "script", "link"]):
        src = tag.get("src") or tag.get("href")
        if not src:
            continue
        if src.startswith(("http://", "https://", "data:")):
            continue

        asset_path = os.path.join(base_path, src.lstrip("/"))
        if not os.path.exists(asset_path):
            missing_assets.append(src)

    if missing_assets:
        results["errors"].append(f"Missing assets: {missing_assets}")

    # ------------------------ Basic Click Area ------------------------
    clickable = (
        soup.find(id="designContainer")
        or soup.find(id="viewport")
        or soup.find(id="clickTagMain")
        or soup.find(id="clickLayer")
        or soup.find(id="clickable")
        or soup.find(attrs={"class": "clickTag"})
        or soup.find(attrs={"class": "clickable"})
    )

    # Extra: detect JS-style anchor click to clickTag
    if not clickable:
        for a in soup.find_all("a"):
            href = a.get("href") or ""
            if "window.open(window.clickTag" in href or "clickTag" in href:
                clickable = a
                break

    if not clickable:
        results["warnings"].append(
            "No clickable area detected in HTML (no click layer or obvious click anchor)."
        )

    return results


# -------------------------------------------------------
# JS VALIDATION
# -------------------------------------------------------
def validate_js(file_path: str) -> dict:
    results = {"warnings": [], "errors": [], "duration": 0, "isi_duration": 0}

    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            js = f.read()
    except Exception:
        results["errors"].append("Error reading JS file.")
        return results

    # Strip single-line comments for easier regex
    js = re.sub(r"//.*", "", js)

    # ------------------------ JS-click detection ------------------------
    click_patterns = [
        r"\.addEventListener\(['\"]click['\"]",
        r"window\.onclick",
        r"document\.onclick",
        r"onclick\s*=",
        r"Enabler\.exit",
        r"\bexit\(",
        r"\bclickTag\b",
    ]
    js_click_found = any(re.search(p, js) for p in click_patterns)
    if js_click_found:
        results["warnings"].append("JS-only click handler detected.")

    # ------------------------ Canvas Border ------------------------
    canvas_patterns = [
        r"strokeRect\s*\(",
        r"lineWidth\s*=\s*1",
        r"strokeStyle\s*=\s*['\"]#?000",
    ]
    for p in canvas_patterns:
        if re.search(p, js):
            results["warnings"].append("Canvas border detected via JS.")
            break

    # ------------------------ Animation Duration ------------------------
    delays = [float(d) for d in re.findall(r"delayedCall\(\s*(\d+\.?\d*)", js)]
    frames = [float(d) for d in re.findall(r"frameDelay\s*=\s*(\d+\.?\d*)", js)]
    animation_duration = sum(delays) or sum(frames)

    isi_scroll = 0.0
    if "scrollTo" in js or "ISIscroll" in js:
        m = re.search(r"scrollSpeed\s*=\s*(\d+\.?\d*)", js)
        isi_scroll = float(m.group(1)) if m else 10.0

    total = animation_duration + isi_scroll

    results["duration"] = round(animation_duration, 1)
    results["isi_duration"] = round(isi_scroll, 1)

    if total > ANIMATION_MAX_DURATION:
        results["errors"].append(
            f"Animation exceeds limit: {total:.1f}s (max {ANIMATION_MAX_DURATION}s)"
        )
    elif total >= ANIMATION_WARNING_THRESHOLD:
        results["warnings"].append(
            f"Animation approaching limit: {total:.1f}s (max {ANIMATION_MAX_DURATION}s)"
        )

    # ------------------------ Loop Count ------------------------
    loops = re.findall(r"repeat\s*:\s*(\d+|Infinity|-1)", js)
    for loop in loops:
        if loop in ("Infinity", "-1") or (loop.isdigit() and int(loop) > MAX_LOOP_COUNT):
            results["errors"].append(f"Loop count exceeds limit: {loop}")

    return results


# -------------------------------------------------------
# Preview Route
# -------------------------------------------------------
@app.route("/preview/<path:folder>/<filename>")
def preview_banner(folder: str, filename: str):
    """
    Serve a specific HTML file for preview.
    folder == "_root_" means files directly under UPLOAD_FOLDER.
    """
    if "__MACOSX" in folder or filename.startswith("."):
        return "Invalid preview request", 404

    if folder == "_root_":
        folder_path = UPLOAD_FOLDER
    else:
        folder_path = os.path.join(UPLOAD_FOLDER, folder)

    full = os.path.join(folder_path, filename)

    if not os.path.exists(full):
        return "Not found", 404

    return send_from_directory(folder_path, filename)


# -------------------------------------------------------
# Index Route
# -------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


# -------------------------------------------------------
# Validation Endpoint (main)
# -------------------------------------------------------
@app.route("/validate", methods=["POST"])
def validate_file():
    cleanup_old_uploads()  # Auto cleanup

    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type"}), 400

    filename = secure_filename(file.filename)
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(file_path)

    file_size_kb = os.path.getsize(file_path) / 1024
    preview_links: list[str] = []
    validation: dict = {"html": {}, "js": {}}
    durations = {"animation": 0.0, "isi": 0.0, "total": 0.0}

    # ------------------------ ZIP Handling ------------------------
    if filename.lower().endswith(".zip"):
        extracted = extract_zip(file_path)
        if not extracted:
            return jsonify({"error": "Invalid ZIP file"}), 400

        html_files: list[str] = []
        js_files: list[str] = []

        for root, _, files in os.walk(extracted):
            if "__MACOSX" in root:
                continue
            for f in files:
                fp = os.path.join(root, f)
                if f.lower().endswith(".html"):
                    html_files.append(fp)
                elif f.lower().endswith(".js"):
                    js_files.append(fp)

        if not html_files:
            return jsonify({"error": 'No HTML file found in ZIP.'}), 400

        # HTML validation + preview link per HTML
        for html in html_files:
            html_name = os.path.basename(html)
            validation["html"][html_name] = validate_html(html)

            rel_folder = os.path.relpath(os.path.dirname(html), UPLOAD_FOLDER)
            preview_url = url_for(
                "preview_banner", folder=rel_folder, filename=html_name
            )
            preview_links.append(preview_url)

        # JS validation
        for js in js_files:
            js_name = os.path.basename(js)
            res = validate_js(js)
            validation["js"][js_name] = res
            durations["animation"] += res["duration"]
            durations["isi"] += res["isi_duration"]

    # ------------------------ Single HTML Upload ------------------------
    elif filename.lower().endswith(".html"):
        validation["html"][filename] = validate_html(file_path)
        preview_url = url_for("preview_banner", folder="_root_", filename=filename)
        preview_links.append(preview_url)

    durations["animation"] = round(durations["animation"], 1)
    durations["isi"] = round(durations["isi"], 1)
    durations["total"] = round(durations["animation"] + durations["isi"], 1)

    return jsonify(
        {
            "validation_results": validation,
            "previews": preview_links,
            "file_size_kb": round(file_size_kb, 2),
            "durations": durations,
        }
    )


# -------------------------------------------------------
# Word Report Endpoint
# -------------------------------------------------------
@app.route("/word-report", methods=["POST"])
def word_report():
    """
    Expect JSON payload like:
    {
      "file_name": "...",
      "file_size_kb": 123.45,
      "validation_results": {...},
      "durations": {...}
    }
    """
    payload = request.get_json()
    if not payload:
        return jsonify({"error": "No data provided"}), 400

    file_name = payload.get("file_name", "Unknown file")
    size_kb = payload.get("file_size_kb")
    validation = payload.get("validation_results", {})
    durations = payload.get("durations", {})

    html_results = validation.get("html", {}) or {}
    js_results = validation.get("js", {}) or {}

    # ---- Summary calculations ----
    missing_assets = any(
        any("Missing assets" in e for e in (res.get("errors") or []))
        for res in html_results.values()
    )

    banner_sizes = [
        res.get("banner_size")
        for res in html_results.values()
        if res.get("banner_size")
    ]
    first_html = next(iter(html_results.values()), None)

    loop_issue = any(
        any("Loop count exceeds" in e for e in (res.get("errors") or []))
        for res in js_results.values()
    )

    clickable_warning = any(
        any("No clickable area" in w for w in (res.get("warnings") or []))
        for res in html_results.values()
    )

    # ---- Build DOCX ----
    doc = Document()
    doc.add_heading("HTML5 Ad Validator Report", level=1)

    p = doc.add_paragraph()
    p.add_run("File: ").bold = True
    p.add_run(file_name)

    if size_kb is not None:
        p = doc.add_paragraph()
        p.add_run("File Size: ").bold = True
        p.add_run(f"{size_kb} KB")

    # Summary lines – flexible style (Option B)
    doc.add_paragraph(
        "Assets Check: "
        + ("✅ All assets present" if not missing_assets else "❌ Missing assets – please review.")
    )

    if banner_sizes:
        doc.add_paragraph("Banner Size(s): " + ", ".join(banner_sizes))

    if first_html:
        doc.add_paragraph("Border Check: " + first_html.get("border", ""))

    doc.add_paragraph(
        "HTML Validation: " + ("✅ Checked" if html_results else "❌ No HTML files found.")
    )

    doc.add_paragraph(
        "Loop Count: "
        + (
            "✅ Within limit"
            if not loop_issue
            else "❌ Loop count exceeds recommended limit – please review."
        )
    )

    doc.add_paragraph(
        "Clickable Area: "
        + ("HTML/JS click detected or not flagged" if not clickable_warning
           else "No clickable area detected – please check manually.")
    )

    doc.add_paragraph(f"Animation Duration: {durations.get('animation', 0)}s")
    doc.add_paragraph(f"ISI Scroll Duration: {durations.get('isi', 0)}s")
    doc.add_paragraph(f"Total Estimated Duration: {durations.get('total', 0)}s")

    # ---- Detailed sections ----
    if html_results:
        doc.add_heading("HTML Details", level=2)
        for name, res in html_results.items():
            doc.add_heading(name, level=3)
            if res.get("banner_size"):
                doc.add_paragraph(f"Banner Size: {res['banner_size']}")
            doc.add_paragraph(f"Border: {res.get('border', '')}")
            if res.get("errors"):
                doc.add_paragraph("Errors:")
                for err in res["errors"]:
                    doc.add_paragraph(err, style="List Bullet")
            if res.get("warnings"):
                doc.add_paragraph("Warnings:")
                for warn in res["warnings"]:
                    doc.add_paragraph(warn, style="List Bullet")

    if js_results:
        doc.add_heading("JS Details", level=2)
        for name, res in js_results.items():
            doc.add_heading(name, level=3)
            doc.add_paragraph(f"Animation Duration: {res.get('duration', 0)}s")
            doc.add_paragraph(f"ISI Scroll Duration: {res.get('isi_duration', 0)}s")
            if res.get("errors"):
                doc.add_paragraph("Errors:")
                for err in res["errors"]:
                    doc.add_paragraph(err, style="List Bullet")
            if res.get("warnings"):
                doc.add_paragraph("Warnings:")
                for warn in res["warnings"]:
                    doc.add_paragraph(warn, style="List Bullet")

    # ---- Return as download ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = os.path.splitext(file_name)[0] or "ad_validation_report"

    return send_file(
        buf,
        as_attachment=True,
        download_name=f"{safe_name}.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


# -------------------------------------------------------
# Run Server
# -------------------------------------------------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    # On Fly this is ignored (they run via entrypoint), but fine locally.
    app.run(host="0.0.0.0", port=port)
